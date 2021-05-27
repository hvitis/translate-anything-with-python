import sys
import json
import requests
import os
import os.path

import pytesseract

from docx import Document
from docx.shared import Pt  # Point font size

try:
    from PIL import Image
except ImportError:
    import Image


# example token
GOOGLE_API_TOKEN = "ya29.c.Kp8BAQg...-e7w1JR786rWdHlSZPzve"

DIR = input("Enter name of the folder with images of scanned book: ")

chosen_path = os.path.join(DIR)

if not os.path.isdir(chosen_path):
    sys.exit("No folder named {}, exiting.".format(DIR))

LANGUAGE_CODES = [
    "af",
    "ga",
    "sq",
    "it",
    "ar",
    "ja",
    "az",
    "kn",
    "eu",
    "ko",
    "bn",
    "la",
    "be",
    "lv",
    "bg",
    "lt",
    "ca",
    "mk",
    "zh-CN",
    "ms",
    "zh-TW",
    "mt",
    "hr",
    "no",
    "cs",
    "fa",
    "da",
    "pl",
    "nl",
    "pt",
    "en",
    "ro",
    "eo",
    "ru",
    "et",
    "sr",
    "tl",
    "sk",
    "fi",
    "sl",
    "fr",
    "es",
    "gl",
    "sw",
    "ka",
    "sv",
    "de",
    "ta",
    "el",
    "te",
    "gu",
    "th",
    "ht",
    "tr",
    "iw",
    "uk",
    "hi",
    "ur",
    "hu",
    "vi",
    "is",
    "cy",
    "id",
    "yi",
]

print(
    """
    Available languages and their Language Codes:
    Language Name	Language Code
    Afrikaans	af
    Irish	ga
    Albanian	sq
    Italian	it
    Arabic	ar
    Japanese	ja
    Azerbaijani	az
    Kannada	kn
    Basque	eu
    Korean	ko
    Bengali	bn
    Latin	la
    Belarusian	be
    Latvian	lv
    Bulgarian	bg
    Lithuanian	lt
    Catalan	ca
    Macedonian	mk
    Chinese Simplified	zh-CN
    Malay	ms
    Chinese Traditional	zh-TW
    Maltese	mt
    Croatian	hr
    Norwegian	no
    Czech	cs
    Persian	fa
    Danish	da
    Polish	pl
    Dutch	nl
    Portuguese	pt
    English	en
    Romanian	ro
    Esperanto	eo
    Russian	ru
    Estonian	et
    Serbian	sr
    Filipino	tl
    Slovak	sk
    Finnish	fi
    Slovenian	sl
    French	fr
    Spanish	es
    Galician	gl
    Swahili	sw
    Georgian	ka
    Swedish	sv
    German	de
    Tamil	ta
    Greek	el
    Telugu	te
    Gujarati	gu
    Thai	th
    Haitian Creole	ht
    Turkish	tr
    Hebrew	iw
    Ukrainian	uk
    Hindi	hi
    Urdu	ur
    Hungarian	hu
    Vietnamese	vi
    Icelandic	is
    Welsh	cy
    Indonesian	id
    Yiddish	yi
    
    Above available languages and their Language Codes
    """
)

LANGUAGE_TO_TRANSLATE = input(
    "Enter Language Code of the language you want to translate to: "
)

if LANGUAGE_TO_TRANSLATE not in LANGUAGE_CODES:
    sys.exit("Chosen Language Code not available, exiting.")


def valid_xml_char_ordinal(c):
    # Cleaning strings for clean save with python-docx
    # Courtesy of StackOverflow's users
    # https://stackoverflow.com/a/8735509/5768457
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF
        or codepoint in (0x9, 0xA, 0xD)
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )


def translate_with_google_api(array_of_original_pages, google_api_token, to):
    # Returns array of strings where each string is a translated text from a scanned page
    print("Translating document to {}...".format(to))
    url = "https://translation.googleapis.com/language/translate/v2"

    headers = {
        "Authorization": "Bearer {}".format(google_api_token),
        "Content-Type": "application/json",
    }
    translation = []
    for page_to_translate in array_of_original_pages:
        print("...")
        payload = json.dumps({"q": [page_to_translate], "target": to})
        response = requests.request("POST", url, headers=headers, data=payload)
        translated_page = json.loads(response.text)
        translated_page = translated_page["data"]["translations"][0]["translatedText"]
        translation.append(translated_page)
    return translation


def ocr_from_images(dir):
    # Returns array of strings where each string is an original text from a scanned page

    files = [
        name for name in os.listdir(dir) if os.path.isfile(os.path.join(dir, name))
    ]
    files.sort()
    array_of_pages = []
    print("Scanning {} images for text...".format(str(len(files))))
    for index, PHOTO in enumerate(files):
        if PHOTO != ".DS_Store":  # for OSX Users
            # If you don't have tesseract executable in your PATH, include the following:
            # pytesseract.pytesseract.tesseract_cmd = r'<full_path_to_your_tesseract_executable>'
            # Example tesseract_cmd = r'C:\Program Files (x86)\Tesseract-OCR\tesseract'
            print("...")
            # Scan the photo to obtain text
            original_text = pytesseract.image_to_string(
                Image.open("{}/{}".format(dir, PHOTO))
            )
            array_of_pages.append(original_text)
    return array_of_pages


def generate_doc(title, array_of_pages):
    # Write the strings from the list into a *.docx file using pip python-docx
    print("Generating document: {}".format(title))
    document = Document()
    document.add_heading(title, 0)

    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(15)

    for index, page_content in enumerate(array_of_pages):
        document.add_heading("Page {}".format(index + 1), level=2)
        cleaned_string = "".join(
            c for c in page_content if valid_xml_char_ordinal(c)
        )  # Cleaning invalid characters generated by pytesseract
        
        paragraph = document.add_paragraph(cleaned_string)
        paragraph.style = style
        document.add_page_break()

    document.save("{}.docx".format(title))


def get_book_title(dir):
    files = [
        name for name in os.listdir(dir) if os.path.isfile(os.path.join(dir, name))
    ]
    book_title = files[0]
    if book_title == ".DS_Store":
        book_title = files[1]
    return book_title.split(".")[0]


original_pages = ocr_from_images(DIR)
generate_doc("{}".format(get_book_title(DIR)), original_pages)

translated_pages = translate_with_google_api(
    original_pages, GOOGLE_API_TOKEN, to=LANGUAGE_TO_TRANSLATE
)
generate_doc("{}_translated".format(get_book_title(DIR)), translated_pages)